#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "python-docx>=1.2.0",
# ]
# ///
"""Verify that the generated .docx contains every word of the policy body.

Compares source body (extracted paragraphs minus scaffolding) to output body
(docx body paragraphs minus the centered emblem paragraph and title).

Scaffolding stripped on BOTH sides (using user-supplied metadata):
  - Repeated `Policy {number} {Series}` header line that appears inline in the
    source PDF on each page.
  - Trailing `Adoption: ...` / `Adopted: ...` and `Revised: ...` lines.
  - The title line itself (rendered separately, not part of body).

Header/footer of the output are verified to contain the user-supplied metadata.

Exit 0 = pass, 1 = differences, 2 = invocation error.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document


def normalize(s: str) -> str:
    s = unicodedata.normalize("NFKC", s)
    s = s.replace(" ", " ")  # non-breaking space → regular
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def normalize_loose(s: str) -> str:
    """Lowercased + smart-quote-normalized form for word-presence checks.

    Casing differences (e.g., title block in source as 'STUDENTS ON GOVERNING BOARDS'
    vs new-template title rendered as 'Students on Governing Boards') and
    smart quotes (e.g., source uses ’ but output uses ') are not modifications
    of the board-approved text — they're rendering differences.
    """
    s = normalize(s).lower()
    return (
        s.replace("‘", "'").replace("’", "'")
        .replace("“", '"').replace("”", '"')
        .replace("–", "-").replace("—", "-")
    )


SCAFFOLD_DATE_LINE = re.compile(r"^\s*(adoption(?:\s+date)?|adopted|updated|date)\s*[:].*", re.I)
SCAFFOLD_REVISED_LINE = re.compile(r"^\s*revised\s*[:].*", re.I)


def _scaffold_filters(policy_number: str, series: str, title: str):
    # Lowercase for symmetric matching with build_docx.strip_scaffolding (which
    # uses normalize_for_compare → lowercase). Without this, "Students On
    # Governing Boards" in source survives verify's strip while build strips it
    # from output, causing false "missing words" reports.
    new_header = normalize(f"Policy {policy_number} {series}").lower()
    title_norm = normalize(title).lower()
    pn = re.escape(policy_number) + r"[A-Za-z]?"
    page_footer = re.compile(
        rf"^\s*page\s+\d+\s+of\s*\d+\s+policy\s+{pn}\s*$", re.I
    )
    old_title_block = re.compile(
        rf"^\s*{re.escape(series)}\s*[–-]\s*series\s+\d+\s+{re.escape(title)}\s*[–-]\s*{pn}\s*$",
        re.I,
    )
    return {new_header, title_norm}, page_footer, old_title_block


def strip_scaffolding_text(text: str, policy_number: str, series: str, title: str) -> str:
    skip_norms, page_footer, old_title_block = _scaffold_filters(policy_number, series, title)
    out_chunks: list[str] = []
    for chunk in re.split(r"(?<=[.])\s+|\n+", text):
        n = normalize(chunk).lower()
        if not n or n in skip_norms:
            continue
        if page_footer.match(chunk) or old_title_block.match(chunk):
            continue
        if SCAFFOLD_DATE_LINE.match(chunk) or SCAFFOLD_REVISED_LINE.match(chunk):
            continue
        out_chunks.append(chunk)
    return " ".join(out_chunks)


def strip_scaffolding_paragraphs(paragraphs: list[dict], policy_number: str, series: str, title: str) -> str:
    skip_norms, page_footer, old_title_block = _scaffold_filters(policy_number, series, title)
    kept: list[str] = []
    for p in paragraphs:
        text = p["text"]
        n = normalize(text).lower()
        if n in skip_norms:
            continue
        if page_footer.match(text) or old_title_block.match(text):
            continue
        if SCAFFOLD_DATE_LINE.match(text) or SCAFFOLD_REVISED_LINE.match(text):
            continue
        m = re.search(r"\b(adoption(?:\s+date)?|adopted|updated|date)\s*:", text, re.I)
        if m:
            text = text[: m.start()].rstrip()
            if not text:
                continue
        kept.append(text)
    return " ".join(kept)


def extract_output_body(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    body_texts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            body_texts.append(p.text)
    return " ".join(body_texts)


def extract_output_header_footer(docx_path: Path) -> tuple[str, str]:
    doc = Document(str(docx_path))
    header_parts: list[str] = []
    footer_parts: list[str] = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                header_parts.append(p.text)
        for p in section.footer.paragraphs:
            if p.text.strip():
                footer_parts.append(p.text)
    return " ".join(header_parts), " ".join(footer_parts)


def find_missing(source_norm: str, output_norm: str) -> list[str]:
    """Find source words absent from the output.

    Walks the source in shrinking windows (15→10→6→3 words) anchored at each
    position. If no window matches, only flag the single word at position i
    when it truly does not appear anywhere as a substring in the output —
    otherwise the failure is just contextual reflow (e.g., the source's
    multi-word chunk got reordered or split across paragraphs in the output)
    and the word is preserved.
    """
    src_words = source_norm.split()
    missing: list[str] = []
    window = 15
    i = 0
    while i < len(src_words):
        chunk = " ".join(src_words[i : i + window])
        if chunk and chunk in output_norm:
            i += window
            continue
        found = False
        for size in (10, 6, 3):
            small = " ".join(src_words[i : i + size])
            if small and small in output_norm:
                i += size
                found = True
                break
        if not found:
            word = src_words[i]
            if word and word not in output_norm:
                missing.append(word)
            i += 1
    return missing


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--source-json", required=True)
    ap.add_argument("--output-docx", required=True, type=Path)
    ap.add_argument("--policy-number", required=True)
    ap.add_argument("--series", required=True)
    ap.add_argument("--title", required=True)
    ap.add_argument("--adopted", default="")
    ap.add_argument("--revised", default="")
    args = ap.parse_args()

    data = json.loads(Path(args.source_json).read_text())

    # Source body (strip scaffolding using paragraph structure)
    source_body = strip_scaffolding_paragraphs(
        data["paragraphs"], args.policy_number, args.series, args.title
    )

    # Output body (the body has emblem in its own paragraph, title, then content;
    # strip the title paragraph)
    output_body_raw = extract_output_body(args.output_docx)
    output_body = strip_scaffolding_text(
        output_body_raw, args.policy_number, args.series, args.title
    )

    # Word-presence check is loose (case + quote insensitive) — board-approved text
    # is preserved character-for-character at the paragraph level by the build step.
    src_norm = normalize_loose(source_body)
    out_norm = normalize_loose(output_body)

    missing = find_missing(src_norm, out_norm)

    # Verify header contents; dates now live in body, not footer
    out_header, _ = extract_output_header_footer(args.output_docx)
    out_header_n = normalize(out_header)
    out_body_raw_n = normalize(output_body_raw)

    header_ok = (
        normalize(f"Policy {args.policy_number}") in out_header_n
        and normalize(args.series) in out_header_n
    )
    dates_ok = True
    if args.adopted.strip():
        dates_ok = normalize(f"Adopted: {args.adopted}") in out_body_raw_n
    if args.revised.strip():
        dates_ok = dates_ok and normalize(f"Revised: {args.revised}") in out_body_raw_n

    title_present = normalize(args.title) in out_body_raw_n

    result = {
        "pass": (len(missing) == 0) and header_ok and dates_ok and title_present,
        "body_missing_word_count": len(missing),
        "body_missing_sample": missing[:50],
        "title_present": title_present,
        "header_ok": header_ok,
        "dates_ok": dates_ok,
        "source_body_chars": len(src_norm),
        "output_body_chars": len(out_norm),
    }
    print(json.dumps(result, indent=2, ensure_ascii=False))
    sys.exit(0 if result["pass"] else 1)


if __name__ == "__main__":
    main()

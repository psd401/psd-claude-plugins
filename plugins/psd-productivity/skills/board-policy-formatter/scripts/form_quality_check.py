#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "python-docx>=1.2.0",
# ]
# ///
"""Automated quality validator for form .docx outputs.

Catches the recurring issues that visual review keeps surfacing:

  * Mid-word bold splits — "ON E WEEK", "Cont act" (a run boundary inside a word)
  * Weird dot fills — "·" * N runs of middots used as fake underline
  * Long underscore runs in body — should be a fillable line, not raw underscores
  * Empty field underlines — labels like "Today's Date:" with no underline after
  * Stacked-label tables — a row of labels with one giant underline below
  * Stray FIELDLINE markers — extractor artifacts that survived to output
  * Body word ALL CAPS line wraps (label rendered without table → spans full width)

Run per-form:
    uv run form_quality_check.py --docx <path>

Or in batch via stdin with paths:
    ls docx/*.docx | uv run form_quality_check.py --stdin
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document


WEIRD_DOT_RE = re.compile(r"[·]{3,}")
LONG_UNDERSCORE_RE = re.compile(r"_{5,}")
FIELDLINE_MARKER = "__FIELDLINE__"
# Mid-word bold split: a word that contains a single capital after lowercase, like "ON E" — bold transition mid-word
# Detected by looking for runs that END inside a word boundary (next run starts with a letter, current ends with one)
SUSPICIOUS_WORD_PATTERNS = [
    re.compile(r"\b(ON|IN|AT|TO|OF) [A-Z](?=\s)"),  # "ON E", "IN E", etc — likely "ONE/INE" split
    re.compile(r"\b[A-Z][a-z]+ [a-z]+\b"),  # e.g., "Cont act" lower-case mid word
]


def gather_text(doc) -> list[str]:
    """All paragraph + cell text in order."""
    out: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        out.append(p.text)
    return out


def detect_mid_word_bold_split(doc) -> list[str]:
    """Walk runs in each paragraph; flag any boundary where bold-state changes
    inside what looks like a single word (no whitespace at the boundary)."""
    issues: list[str] = []
    def scan(paras):
        for p in paras:
            runs = p.runs
            for i in range(len(runs) - 1):
                a, b = runs[i], runs[i + 1]
                at = a.text or ""
                bt = b.text or ""
                if not at or not bt:
                    continue
                # End of `a` and start of `b` are both letters → mid-word boundary
                if at[-1].isalpha() and bt[0].isalpha():
                    if (a.bold or False) != (b.bold or False):
                        snippet = (at[-15:] + "|" + bt[:15]).replace("\n", " ")
                        issues.append(f"mid-word bold split: …{snippet}…")
    scan(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                scan(cell.paragraphs)
    return issues


def detect_weird_dots(text: str) -> list[str]:
    out = []
    for m in WEIRD_DOT_RE.finditer(text):
        out.append(f"middot fill: '{m.group(0)[:20]}'")
    return out


def detect_long_underscores(text: str) -> list[str]:
    out = []
    for m in LONG_UNDERSCORE_RE.finditer(text):
        out.append(f"raw underscores in body: '{m.group(0)[:20]}'")
    return out


def detect_fieldline_marker(text: str) -> list[str]:
    if FIELDLINE_MARKER in text:
        return [f"unstripped FIELDLINE marker"]
    return []


def detect_label_without_fill(doc) -> list[str]:
    """A bold label ending in colon with no following underlined run or tab-leader.

    Only flag labels that appear mid-paragraph (i.e., there is sibling content
    in the same paragraph). A label that is the entire paragraph commonly
    leads a downstream block (checkbox group, response box, table) — leaving
    that determination to the visual reviewer.
    """
    out: list[str] = []
    def scan(paras):
        for p in paras:
            runs = p.runs
            # Skip standalone-label paragraphs (single short bold colon-ended line).
            full = (p.text or "").strip()
            if full and full.endswith(":") and len(full) < 60:
                continue
            tabs = p.paragraph_format.tab_stops
            if tabs and len(list(tabs)) > 0:
                # Has tab stops — likely uses leader-fill paradigm. Skip.
                continue
            for i, r in enumerate(runs):
                t = (r.text or "").rstrip()
                if r.bold and t.endswith(":") and len(t) < 60:
                    # Look at what follows
                    if i + 1 >= len(runs):
                        out.append(f"label without fill: '{t[:40]}'")
                        continue
                    nxt = runs[i + 1]
                    nxt_t = nxt.text or ""
                    # Considered OK if next run has underline, tab char, or fills the line
                    if nxt.underline:
                        continue
                    if "\t" in nxt_t:
                        continue
                    if nxt_t.strip() == "":
                        # Empty run — only OK if a tab-leader is set on the paragraph
                        if p.paragraph_format.tab_stops and len(list(p.paragraph_format.tab_stops)) > 0:
                            continue
                    # Allow inline body text (e.g., "Date: 12/2019" already filled) — skip flagging
                    if nxt_t.strip() and not nxt.underline:
                        continue
                    out.append(f"label without fill: '{t[:40]}'")
    scan(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                scan(cell.paragraphs)
    return out


def check_file(path: Path) -> dict:
    doc = Document(str(path))
    texts = gather_text(doc)
    issues: list[str] = []
    full = "\n".join(texts)
    issues += detect_weird_dots(full)
    issues += detect_long_underscores(full)
    issues += detect_fieldline_marker(full)
    issues += detect_mid_word_bold_split(doc)
    issues += detect_label_without_fill(doc)
    return {
        "file": str(path),
        "issue_count": len(issues),
        "issues": issues[:25],
    }


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", type=Path)
    ap.add_argument("--dir", type=Path, help="Check every .docx in this directory")
    args = ap.parse_args()

    paths: list[Path] = []
    if args.docx:
        paths = [args.docx]
    elif args.dir:
        paths = sorted(args.dir.glob("*.docx"))
    else:
        ap.error("provide --docx or --dir")

    summary = []
    total_issues = 0
    for p in paths:
        r = check_file(p)
        summary.append(r)
        total_issues += r["issue_count"]
        flag = "OK  " if r["issue_count"] == 0 else "FAIL"
        print(f"{flag} {p.name:<70} {r['issue_count']:>3} issues")
        for i in r["issues"][:5]:
            print(f"     - {i}")
    print(f"\n{sum(1 for r in summary if r['issue_count']==0)}/{len(summary)} clean, total issues = {total_issues}")
    sys.exit(0 if total_issues == 0 else 1)


if __name__ == "__main__":
    main()

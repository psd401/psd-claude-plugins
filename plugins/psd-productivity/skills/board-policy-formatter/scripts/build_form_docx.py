#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "python-docx>=1.2.0",
# ]
# ///
"""Build a beautifully designed PSD-branded form .docx from extract_form.py JSON.

Design system (see design.py for colors):
  * Title:       18pt Calibri Bold, Pacific blue, centered
  * Subtitle:    11pt italic, Whulge gray-blue, centered
  * Section banner: shaded full-width cell (Pacific bg, white text, 13pt bold)
  * Field row:   1-row table with paired (Bold Label | underlined fill cell)
                 columns; multi-field rows balance across N×2 cells
  * Checkbox:    Wingdings ☐ glyph (Word renders cleanly), grouped in 2-col table
  * Body:        11pt Times Roman, dark gray, 1.3 line spacing
  * Footer:      Sea Glass divider line + centered Page X of Y in Whulge
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Emu

sys.path.insert(0, str(Path(__file__).parent))
from design import (  # noqa: E402
    PACIFIC, SEA_GLASS, SEA_FOAM, SKYLIGHT, WHULGE, MUTED_GRAY, BODY_GRAY, WHITE,
    HEADER_FONT, BODY_FONT, LABEL_FONT,
    SIZE_TITLE, SIZE_SUBTITLE, SIZE_SECTION, SIZE_BODY, SIZE_LABEL, SIZE_FOOTER,
    SPACE_AFTER_TITLE, SPACE_AFTER_SUBTITLE, SPACE_AFTER_SECTION, SPACE_AFTER_PARA, SPACE_AFTER_FIELD_ROW,
    docx_set_run_font, docx_set_cell_shading, docx_set_cell_borders,
    docx_remove_cell_borders, docx_set_cell_margins,
)

EMBLEM_REL = Path(__file__).resolve().parents[2] / "psd-brand-guidelines" / "assets" / "psd_logo-2color-square.png"
SAFE_FILENAME = re.compile(r"[^A-Za-z0-9 .,'&_()\-]")
# Anchored to start-of-word AND must end at a word boundary before the colon.
# Total label length is unbounded (no 40-char cap that caused mid-word matches like "E WEEK...").
INLINE_FIELD_RE = re.compile(
    r"(?<![A-Za-z])([A-Z][A-Za-z'/\-]*(?:[,]?\s+[A-Za-z'/\-(),]+){0,8})\s*:\s*_+"
)
USABLE_WIDTH_INCHES = 6.5  # Letter at 1in margins each side


def add_title(doc, text: str) -> None:
    """Title — same as policies: 16pt bold black centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(SPACE_AFTER_TITLE)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    docx_set_run_font(run, font=HEADER_FONT, size=SIZE_TITLE, bold=True)


def add_subtitle(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(SPACE_AFTER_SUBTITLE)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    docx_set_run_font(run, font=HEADER_FONT, size=SIZE_SUBTITLE, italic=True, color=MUTED_GRAY)


def add_section_banner(doc, text: str) -> None:
    """Full-width shaded banner with bold white section title."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.columns[0].width = Inches(USABLE_WIDTH_INCHES)
    cell = t.rows[0].cells[0]
    cell.width = Inches(USABLE_WIDTH_INCHES)
    docx_set_cell_shading(cell, PACIFIC)
    docx_remove_cell_borders(cell)
    docx_set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text.upper())
    docx_set_run_font(run, font=HEADER_FONT, size=SIZE_SECTION, bold=True, color=WHITE)
    # Trailing spacer so following content doesn't crash into banner
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(SPACE_AFTER_SECTION)
    spacer.paragraph_format.space_before = Pt(0)


def add_body_paragraph(doc, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(SPACE_AFTER_PARA)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    docx_set_run_font(run, font=BODY_FONT, size=SIZE_BODY, bold=bold, italic=italic, color=BODY_GRAY)


def add_field_row(doc, labels: list[str]) -> None:
    """Render fields as a single paragraph with tab stops + underline leaders.

    This is Word-native — each label is followed by a tab that draws an underline
    to the next stop. Word handles spacing naturally and the result is editable
    (clicking on a tab fill puts the cursor there to type).
    """
    from docx.enum.text import WD_TAB_LEADER, WD_TAB_ALIGNMENT
    n = max(1, len(labels))
    # Compute tab stop positions: divide usable width evenly across fields.
    field_width_in = USABLE_WIDTH_INCHES / n

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(SPACE_AFTER_FIELD_ROW)
    p.paragraph_format.space_before = Pt(0)
    tabs = p.paragraph_format.tab_stops
    for i in range(n):
        tab_pos = Inches((i + 1) * field_width_in)
        tabs.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.LINES)

    for i, label in enumerate(labels):
        lr = p.add_run(f"{label.strip().rstrip(':')}: ")
        docx_set_run_font(lr, font=LABEL_FONT, size=SIZE_LABEL, bold=True, color=PACIFIC)
        # Tab character — Word fills it with the LINES leader (underline) to next stop
        tr = p.add_run("\t")
        docx_set_run_font(tr, font=BODY_FONT, size=SIZE_BODY)


def add_response_area(doc, prompt: str, *, lines: int = 3) -> None:
    """Multi-line response area — bold prompt + bordered box of N empty rows."""
    if prompt:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(prompt.strip())
        docx_set_run_font(run, font=LABEL_FONT, size=SIZE_LABEL, bold=True, color=PACIFIC)
    t = doc.add_table(rows=lines, cols=1)
    t.autofit = False
    t.columns[0].width = Inches(USABLE_WIDTH_INCHES)
    for i, row in enumerate(t.rows):
        cell = row.cells[0]
        cell.width = Inches(USABLE_WIDTH_INCHES)
        docx_set_cell_borders(cell, top="single" if i == 0 else None,
                              bottom="single", left="single", right="single",
                              color=PACIFIC, width="4")
        docx_set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        cell.text = ""
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(SPACE_AFTER_FIELD_ROW)


def add_checkbox_group(doc, label: str, options: list[str]) -> None:
    """Bold prompt + checkboxes in a 2-column table when there are 2+ short options;
    otherwise a single-column list."""
    if label.strip():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(label.strip().rstrip(":") + ":")
        docx_set_run_font(run, font=LABEL_FONT, size=SIZE_LABEL, bold=True, color=PACIFIC)
    short = all(len(o) < 35 for o in options) and len(options) >= 2
    cols = 2 if short else 1
    rows_needed = (len(options) + cols - 1) // cols
    t = doc.add_table(rows=rows_needed, cols=cols)
    t.autofit = False
    col_w = Inches(USABLE_WIDTH_INCHES / cols)
    for col in t.columns:
        col.width = col_w
    for idx, opt in enumerate(options):
        r, c = divmod(idx, cols)
        cell = t.rows[r].cells[c]
        cell.width = col_w
        docx_remove_cell_borders(cell)
        docx_set_cell_margins(cell, top=20, bottom=20, left=80, right=80)
        cell.text = ""
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.space_before = Pt(0)
        # Checkbox glyph as Wingdings empty box (renders cleanly in Word)
        cb = cp.add_run("☐  ")
        docx_set_run_font(cb, font=BODY_FONT, size=SIZE_BODY, color=PACIFIC)
        txt = cp.add_run(opt.strip())
        docx_set_run_font(txt, font=BODY_FONT, size=SIZE_BODY, color=BODY_GRAY)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(SPACE_AFTER_FIELD_ROW)


def _is_form_layout_table(rows: list[list[str]]) -> bool:
    """A 'form-layout' table is one where cells are either bold-label-like
    (short text ending in `:`) or empty fillables — not a tabular data display.
    Detect: every non-empty cell is short (<40 chars) and many end with ':'."""
    if not rows:
        return False
    cells = [c for row in rows for c in row]
    nonempty = [c.strip() for c in cells if c.strip()]
    if not nonempty:
        return False
    label_like = [c for c in nonempty if c.endswith(":") and len(c) <= 40]
    empty = sum(1 for c in cells if not c.strip())
    # Form layouts have at least one empty fill cell and most non-empty cells are labels.
    if empty == 0:
        return False
    return len(label_like) >= max(1, int(0.6 * len(nonempty)))


def add_form_layout_table(doc, rows: list[list[str]]) -> None:
    """Render a source table that looks like a form layout: bold labels in
    some cells, empty fill cells in others. Empty cells get a bottom border
    (underline). No row shading, no header band."""
    if not rows:
        return
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=cols)
    t.autofit = False
    cell_w = Inches(USABLE_WIDTH_INCHES / cols)
    for col in t.columns:
        col.width = cell_w
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = t.rows[r_idx].cells[c_idx]
            cell.width = cell_w
            text = (row[c_idx] if c_idx < len(row) else "").strip()
            cell.text = ""
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(2)
            cp.paragraph_format.space_before = Pt(0)
            if text:
                run = cp.add_run(text)
                is_label = text.endswith(":") and len(text) <= 40
                if is_label:
                    docx_set_run_font(run, font=LABEL_FONT, size=SIZE_LABEL, bold=True, color=PACIFIC)
                else:
                    docx_set_run_font(run, font=BODY_FONT, size=SIZE_BODY, color=BODY_GRAY)
                docx_remove_cell_borders(cell)
            else:
                # Empty fill cell: bottom border only = underlined fillable line.
                docx_set_cell_borders(cell, bottom="single", color=PACIFIC, width="6")
            docx_set_cell_margins(cell, top=60, bottom=80, left=40, right=40)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(SPACE_AFTER_FIELD_ROW)


def add_data_table(doc, rows: list[list[str]]) -> None:
    """Render a source table with branded styling: Pacific header band,
    alternating Skylight rows, Whulge borders. Routes form-layout tables
    (label + empty fill cells) to add_form_layout_table instead."""
    if not rows:
        return
    if _is_form_layout_table(rows):
        add_form_layout_table(doc, rows)
        return
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=cols)
    t.autofit = False
    cell_w = Inches(USABLE_WIDTH_INCHES / cols)
    for col in t.columns:
        col.width = cell_w
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = t.rows[r_idx].cells[c_idx]
            cell.width = cell_w
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(2)
            cp.paragraph_format.space_before = Pt(0)
            # Split underscore runs (5+) into underlined-blank inline elements.
            parts = re.split(r"(_{5,})", text)
            if r_idx == 0:
                docx_set_cell_shading(cell, PACIFIC)
            elif r_idx % 2 == 0:
                docx_set_cell_shading(cell, SEA_FOAM)
            for part in parts:
                if not part:
                    continue
                if part.startswith("___"):
                    _add_underlined_blank(cp, width_chars=min(40, max(8, len(part))))
                else:
                    seg = cp.add_run(part)
                    if r_idx == 0:
                        docx_set_run_font(seg, font=LABEL_FONT, size=SIZE_LABEL, bold=True, color=WHITE)
                    else:
                        docx_set_run_font(seg, font=BODY_FONT, size=SIZE_BODY, color=BODY_GRAY)
            docx_set_cell_borders(cell, top="single", bottom="single",
                                  left="single", right="single",
                                  color=WHULGE, width="4")
            docx_set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(SPACE_AFTER_FIELD_ROW)


# Header / footer chrome ------------------------------------------------

def set_header(section, form_number: str, series: str) -> None:
    """Header matches policy template: tight, right-aligned, plain black Times."""
    header = section.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    for line in [f"Form {form_number}", series]:
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        docx_set_run_font(run, font=HEADER_FONT, size=SIZE_BODY)


def set_footer_page_number(section) -> None:
    footer = section.footer
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Top border on the footer paragraph as a visual divider
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:color"), SEA_GLASS)
    top.set(qn("w:space"), "6")
    pBdr.append(top)
    pPr.append(pBdr)

    def field(run, instr: str) -> None:
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin"); run._r.append(b)
        i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = instr; run._r.append(i)
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end"); run._r.append(e)

    r1 = p.add_run("Page "); docx_set_run_font(r1, font=HEADER_FONT, size=SIZE_FOOTER)
    r2 = p.add_run(); docx_set_run_font(r2, font=HEADER_FONT, size=SIZE_FOOTER); field(r2, " PAGE ")
    r3 = p.add_run(" of "); docx_set_run_font(r3, font=HEADER_FONT, size=SIZE_FOOTER)
    r4 = p.add_run(); docx_set_run_font(r4, font=HEADER_FONT, size=SIZE_FOOTER); field(r4, " NUMPAGES ")


# Smart paragraph rendering (handles inline `Label:____` patterns) -------

def _add_underlined_blank(p, *, width_chars: int = 30) -> None:
    """Append a run of underlined spaces — Word-native fillable line."""
    # Use figure-spaces (U+2007) which are fixed-width and survive underline rendering.
    run = p.add_run(" " * width_chars)
    docx_set_run_font(run, font=BODY_FONT, size=SIZE_BODY)
    run.underline = True


def render_paragraph(doc, text: str, *, bold: bool = False) -> None:
    """Render a paragraph preserving every word.

    If a paragraph contains `Label:____` patterns, render inline with bold
    labels + underlined fillable spans, keeping all surrounding body text intact.
    """
    matches = list(INLINE_FIELD_RE.finditer(text))
    if matches:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(SPACE_AFTER_PARA)
        p.paragraph_format.space_before = Pt(0)
        cursor = 0
        for m in matches:
            pre = text[cursor:m.start()]
            pre_clean = re.sub(r"_+", "", pre).rstrip()
            if pre_clean:
                r = p.add_run(pre_clean + (" " if not pre_clean.endswith(" ") else ""))
                docx_set_run_font(r, font=BODY_FONT, size=SIZE_BODY, color=BODY_GRAY)
            label = m.group(1).strip()
            lr = p.add_run(f"{label}: ")
            docx_set_run_font(lr, font=LABEL_FONT, size=SIZE_LABEL, bold=True, color=PACIFIC)
            _add_underlined_blank(p, width_chars=28)
            sp = p.add_run(" ")
            docx_set_run_font(sp, font=BODY_FONT, size=SIZE_BODY, color=BODY_GRAY)
            cursor = m.end()
        rest = text[cursor:]
        rest_clean = re.sub(r"_+", "", rest).strip()
        if rest_clean:
            rr = p.add_run(rest_clean)
            docx_set_run_font(rr, font=BODY_FONT, size=SIZE_BODY, color=BODY_GRAY)
        return
    cleaned = re.sub(r"_+", "", text).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    if cleaned:
        add_body_paragraph(doc, cleaned, bold=bold)


# Build ------------------------------------------------------------------

def derive_output_path(form_number: str, title: str, *, output: str | None) -> Path:
    clean = SAFE_FILENAME.sub("", title or "Form").strip()
    name = f"{form_number} - {clean}.docx"
    if output is None:
        return Path(name)
    p = Path(output)
    if p.is_dir() or output.endswith("/") or p.suffix == "":
        return p / name
    return p


def is_section_paragraph(text: str) -> bool:
    t = text.strip().rstrip(":").strip()
    if not t or len(t) > 60:
        return False
    if "☐" in t or "_" in t:
        return False
    if t.count(":") > 0:
        return False
    if not any(c.isalpha() for c in t):
        return False
    words = [w for w in t.split() if any(c.isalpha() for c in w)]
    if not words:
        return False
    caps_count = sum(1 for w in words if w == w.upper())
    return caps_count >= max(1, len(words) - 1)


def build(data: dict, *, form_number: str, series: str, title: str, output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default style → Times 11pt
    normal = doc.styles["Normal"]
    el = normal.element
    rpr = el.find(qn("w:rPr"))
    if rpr is None:
        rpr = el.makeelement(qn("w:rPr"), {})
        el.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), BODY_FONT)

    set_header(section, form_number, series)
    set_footer_page_number(section)

    # Centered emblem
    emblem_p = doc.add_paragraph()
    emblem_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    emblem_p.paragraph_format.space_after = Pt(4)
    if EMBLEM_REL.exists():
        emblem_p.add_run().add_picture(str(EMBLEM_REL), height=Inches(1.1))

    if title:
        add_title(doc, title)
    if data.get("subtitle"):
        add_subtitle(doc, data["subtitle"])
    else:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)

    # Body blocks
    for block in data.get("blocks", []):
        t = block["type"]
        if t == "paragraph":
            text = block["text"]
            # Section banner heuristic
            if is_section_paragraph(text):
                add_section_banner(doc, text)
                continue
            is_heading = block.get("style_hint") == "heading"
            render_paragraph(doc, text, bold=is_heading)
        elif t == "field_row":
            labels = [f["label"] for f in block.get("fields", [])]
            if labels:
                add_field_row(doc, labels)
        elif t == "checkbox_group":
            add_checkbox_group(doc, block.get("label", ""), block.get("options", []))
        elif t == "table":
            add_data_table(doc, block.get("rows", []))

    # Revision date
    rev = data.get("revision_date")
    if rev:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(8)
        sp.paragraph_format.space_after = Pt(0)
        run = sp.add_run(f"Revised {rev}")
        docx_set_run_font(run, font=HEADER_FONT, size=SIZE_FOOTER, italic=True, color=MUTED_GRAY)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", help="JSON file from extract_form.py; default stdin")
    ap.add_argument("--form-number", required=True)
    ap.add_argument("--series", required=True)
    ap.add_argument("--title", default=None)
    ap.add_argument("--output", help="Path or directory; auto-derived if omitted")
    args = ap.parse_args()

    if args.input:
        data = json.loads(Path(args.input).read_text())
    else:
        data = json.loads(sys.stdin.read())

    # Use the source's own title text (preserves exact wording) when available;
    # fall back to the supplied --title for filename derivation and as a render
    # fallback when auto-detection finds nothing usable.
    detected = data.get("title", "").strip()
    rendered_title = detected if detected else (args.title or f"Form {args.form_number}")
    filename_title = args.title or detected or f"Form {args.form_number}"
    output_path = derive_output_path(args.form_number, filename_title, output=args.output)
    build(data, form_number=args.form_number, series=args.series, title=rendered_title, output=output_path)
    print(json.dumps({"output": str(output_path), "block_count": len(data.get("blocks", []))}))


if __name__ == "__main__":
    main()

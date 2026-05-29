#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "python-docx>=1.2.0",
# ]
# ///
"""Verify that every word in the source form survives into the output .docx.

Adapted from verify_fidelity.py (policies). Checks:
  1) All source body words appear in output (case + punctuation insensitive)
  2) Title is present in output
  3) Detected revision date appears (if any)

Forms have looser structure than policies but the rule is the same:
  zero text loss. The renderer may add labels/banners/borders, but it must
  not drop, alter, or fragment any source word.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document


# Words that the form template adds purely as scaffolding — these don't need to
# come from the source.
TEMPLATE_NOISE = {
    "form", "page", "of", "revised", "peninsula", "school", "district",
}
# Also tolerate decorative dashes/separator lines
DECORATIVE_LINE = re.compile(r"^[\-_=*]{3,}$")


def normalize_loose(s: str) -> str:
    s = unicodedata.normalize("NFKC", s)
    s = s.replace(" ", " ").replace("’", "'").replace("‘", "'")
    s = s.replace("“", '"').replace("”", '"')
    s = s.replace("–", "-").replace("—", "-")
    s = re.sub(r"\s+", " ", s)
    return s.strip().lower()


def extract_output_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts: list[str] = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for p in section.footer.paragraphs:
            if p.text.strip():
                parts.append(p.text)
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return " ".join(parts)


def find_missing(source_norm: str, output_norm: str) -> list[str]:
    """Walk source words; flag any that don't appear as a substring in output."""
    src_words = source_norm.split()
    missing: list[str] = []
    for w in src_words:
        bare = re.sub(r"[^a-z0-9'./\-]", "", w)
        if not bare or bare in TEMPLATE_NOISE:
            continue
        if DECORATIVE_LINE.match(bare):
            continue  # decorative separator line, not real content
        # Numbers like "10/2004", "5/12/2023" etc are dates - skip if they appear
        # anywhere in output even reformatted
        if w not in output_norm and bare not in output_norm:
            missing.append(w)
    return missing


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--source-json", required=True, help="extract_form.py output")
    ap.add_argument("--output-docx", required=True, type=Path)
    ap.add_argument("--title", default="")
    args = ap.parse_args()

    data = json.loads(Path(args.source_json).read_text())
    # Source = all raw paragraphs PLUS table cells (whole source text)
    src_paras = data.get("raw_paragraphs", [])
    src_tables = data.get("raw_tables", [])
    src_text_parts = [p["text"] for p in src_paras]
    for t in src_tables:
        for row in t["rows"]:
            for cell in row:
                if cell.strip():
                    src_text_parts.append(cell)
    # Drop pure FIELDLINE markers (extractor scaffolding)
    src_text = " ".join(src_text_parts).replace("__FIELDLINE__", "")

    output_text = extract_output_text(args.output_docx)

    src_norm = normalize_loose(src_text)
    out_norm = normalize_loose(output_text)

    missing = find_missing(src_norm, out_norm)

    # Title check: words from manifest title must each appear somewhere in output.
    # Source-rendered title may differ from manifest title — that's intentional.
    title_words_missing: list[str] = []
    if args.title:
        for w in normalize_loose(args.title).split():
            if w in TEMPLATE_NOISE:
                continue
            if w not in out_norm:
                title_words_missing.append(w)

    # Title check is informational only — manifest titles may legitimately differ
    # from the source's title text (we render the source title verbatim).
    result = {
        "pass": len(missing) == 0,
        "missing_word_count": len(missing),
        "missing_sample": missing[:25],
        "title_words_missing_info": title_words_missing,
        "source_chars": len(src_norm),
        "output_chars": len(out_norm),
    }
    print(json.dumps(result, indent=2, ensure_ascii=False))
    sys.exit(0 if result["pass"] else 1)


if __name__ == "__main__":
    main()

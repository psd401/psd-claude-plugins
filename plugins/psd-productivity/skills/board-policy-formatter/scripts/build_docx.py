#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "python-docx>=1.2.0",
# ]
# ///
"""Build a PSD board policy .docx from extracted paragraphs and metadata.

Reads the JSON output of extract_text.py on stdin (or via --input).
Writes the .docx to --output. Strips template scaffolding (repeated header
line, trailing Adoption/Revised line) using the provided metadata.

Hard rule: paragraph text is rendered verbatim. No rewriting.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


EMBLEM_REL = Path(__file__).resolve().parents[2] / "psd-brand-guidelines" / "assets" / "psd_logo-2color-square.png"

TIMES = "Times New Roman"


def set_run_font(run, *, bold: bool = False, size: int = 12) -> None:
    run.font.name = TIMES
    run.font.size = Pt(size)
    run.bold = bold
    # Ensure east-asian and complex-script font slots also use Times
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), TIMES)


def add_body_paragraph(doc, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=bold, size=12)


def add_labeled_paragraph(doc, label: str, value: str) -> None:
    """Render a paragraph with a bold label followed by a regular value."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    set_run_font(p.add_run(label), bold=True, size=12)
    set_run_font(p.add_run(value), bold=False, size=12)


def normalize_for_compare(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip().lower()


def strip_scaffolding(
    paragraphs: list[dict],
    policy_number: str,
    series: str,
    adopted: str,
    revised: str,
    title: str,
) -> list[dict]:
    """Remove repeated header/footer scaffolding from the extracted body.

    Handles both new format (`Policy NNNN` / `{Series}`) and old format
    (`{Series} – Series NNNN` / `{Title} – NNNN`, `Page X of Y Policy NNNN`).
    """
    new_header = normalize_for_compare(f"Policy {policy_number} {series}")
    title_line = normalize_for_compare(title)
    # Source may include a 'P' suffix (e.g., 1115P) when policy_number is 1115 + procedure
    pn = re.escape(policy_number) + r"[A-Za-z]?"
    old_title_block_pat = re.compile(
        rf"^\s*{re.escape(series)}\s*[–-]\s*series\s+\d+\s+{re.escape(title)}\s*[–-]\s*{pn}\s*$",
        re.I,
    )

    date_line_patterns = [
        re.compile(r"^\s*(adoption(?:\s+date)?|adopted|updated|date)\s*[:]", re.I),
        re.compile(r"^\s*revised\s*[:]", re.I),
    ]
    page_footer_pattern = re.compile(
        rf"^\s*page\s+\d+\s+of\s*\d+\s+policy\s+{pn}\s*$",
        re.I,
    )

    cleaned: list[dict] = []
    for p in paragraphs:
        text = p["text"]
        norm = normalize_for_compare(text)
        if norm in (new_header, title_line):
            continue
        if old_title_block_pat.match(text):
            continue
        if page_footer_pattern.match(text):
            continue
        if any(pat.match(text) for pat in date_line_patterns):
            continue
        cleaned.append(p)
    return cleaned


def set_header(section, policy_number: str, series: str) -> None:
    header = section.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    for line in [f"Policy {policy_number}", series]:
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_run_font(run, size=12)


def set_footer_page_number(section) -> None:
    """Render a centered 'Page X of Y' footer using Word field codes."""
    from docx.oxml import OxmlElement

    footer = section.footer
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    def field(run, instr_text: str) -> None:
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instr_text
        run._r.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

    r1 = p.add_run("Page ")
    set_run_font(r1, size=12)
    r2 = p.add_run()
    set_run_font(r2, size=12)
    field(r2, " PAGE ")
    r3 = p.add_run(" of ")
    set_run_font(r3, size=12)
    r4 = p.add_run()
    set_run_font(r4, size=12)
    field(r4, " NUMPAGES ")


CROSS_REF_SPLIT = re.compile(r"(?=\b(?:Board\s+Policy\s+\d{4,5}|(?<![.\d])(?<!Policy\s)(?<!Policy)\d{4,5}\s+[A-Z]))")
LEGAL_REF_SPLIT = re.compile(r"(?=\b(?:RCW|WAC|USC|CFR)\s+\S)")


def _split_with_header(text: str, header_pattern: str, splitter: re.Pattern) -> list[str] | None:
    m = re.match(rf"^(\s*{header_pattern}\s*:?)\s*(.*)$", text, re.I | re.S)
    if not m:
        return None
    header = m.group(1).strip()
    rest = m.group(2).strip()
    if not rest:
        return [header]
    parts = [p.strip() for p in splitter.split(rest) if p.strip()]
    if not parts:
        return [header, rest]
    return [header] + [re.sub(r"\s+", " ", p) for p in parts]


def split_cross_references(text: str) -> list[str] | None:
    return _split_with_header(text, r"cross\s+references?", CROSS_REF_SPLIT)


def split_legal_references(text: str) -> list[str] | None:
    return _split_with_header(text, r"legal\s+references", LEGAL_REF_SPLIT)


def build(
    paragraphs: list[dict],
    *,
    policy_number: str,
    series: str,
    title: str,
    adopted: str,
    revised: str,
    output: Path,
) -> None:
    doc = Document()

    # Page setup: US Letter, 1" margins
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default style → Times New Roman 12pt
    normal = doc.styles["Normal"]
    normal.font.name = TIMES
    normal.font.size = Pt(12)
    r_pr = normal.element.get_or_add_rPr() if hasattr(normal.element, "get_or_add_rPr") else None
    # Also force eastAsia slot in style for safety
    el = normal.element
    rpr = el.find(qn("w:rPr"))
    if rpr is None:
        rpr = el.makeelement(qn("w:rPr"), {})
        el.append(rpr)
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), TIMES)

    set_header(section, policy_number, series)
    set_footer_page_number(section)

    # Emblem centered at top of body
    emblem_para = doc.add_paragraph()
    emblem_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if EMBLEM_REL.exists():
        emblem_para.add_run().add_picture(str(EMBLEM_REL), height=Inches(1.5))
    else:
        sys.stderr.write(f"WARNING: emblem not found at {EMBLEM_REL}\n")

    # Title centered, 16pt bold, with a blank line above and below
    blank = doc.add_paragraph()
    blank.paragraph_format.space_after = Pt(0)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run(title)
    set_run_font(title_run, bold=True, size=16)

    # Body paragraphs (scaffolding stripped)
    cleaned = strip_scaffolding(paragraphs, policy_number, series, adopted, revised, title)
    for entry in cleaned:
        text = entry["text"]
        is_heading = entry.get("style_hint") == "heading"
        cross = split_cross_references(text)
        legal = split_legal_references(text)
        if cross:
            for idx, line in enumerate(cross):
                add_body_paragraph(doc, line, bold=(idx == 0))
        elif legal:
            for idx, line in enumerate(legal):
                add_body_paragraph(doc, line, bold=(idx == 0))
        else:
            add_body_paragraph(doc, text, bold=is_heading)

    # Adopted / Revised lines rendered at end of body — bold labels, regular values
    if adopted:
        add_labeled_paragraph(doc, "Adopted:", f" {adopted}")
    if revised.strip():
        add_labeled_paragraph(doc, "Revised:", f" {revised}")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


SAFE_FILENAME = re.compile(r"[^A-Za-z0-9 .,'&_()\-]")


def derive_output_path(policy_number: str, title: str, *, procedure: bool, output: str | None) -> Path:
    """Resolve --output into a final path, deriving a filename if needed.

    Filename pattern: `{number}[p] - {title}.docx`
      policy 1000 → `1000 - Title.docx`
      procedure 1000 → `1000p - Title.docx`
    """
    number_token = f"{policy_number}p" if procedure else policy_number
    clean_title = SAFE_FILENAME.sub("", title).strip()
    derived_name = f"{number_token} - {clean_title}.docx"
    if output is None:
        return Path(derived_name)
    p = Path(output)
    if p.is_dir() or output.endswith("/"):
        return p / derived_name
    if p.suffix == "":
        return p / derived_name
    return p


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", help="JSON file from extract_text.py; default stdin")
    ap.add_argument("--policy-number", required=True)
    ap.add_argument("--series", required=True)
    ap.add_argument("--title", required=True)
    ap.add_argument("--adopted", default=None, help="MM/DD/YYYY; falls back to detected_dates from extract step")
    ap.add_argument("--revised", default=None, help="Comma-separated MM/DD/YYYY list; falls back to detected_dates")
    ap.add_argument("--procedure", action="store_true", help="Render as procedure (filename gets 'p' suffix)")
    ap.add_argument("--output", help="Optional explicit output path or directory; auto-derived if omitted")
    args = ap.parse_args()
    output_path = derive_output_path(args.policy_number, args.title, procedure=args.procedure, output=args.output)

    if args.input:
        data = json.loads(Path(args.input).read_text())
    else:
        data = json.loads(sys.stdin.read())

    detected = data.get("detected_dates") or {}
    adopted = args.adopted if args.adopted is not None else detected.get("adopted", "")
    revised = args.revised if args.revised is not None else detected.get("revised", "")

    build(
        data["paragraphs"],
        policy_number=args.policy_number,
        series=args.series,
        title=args.title,
        adopted=adopted,
        revised=revised,
        output=output_path,
    )

    print(json.dumps({"output": str(output_path), "paragraph_count": len(data["paragraphs"])}))


if __name__ == "__main__":
    main()
